"""Convert numbskull spud markdown to a styled Word document."""
import re
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import latex2mathml.converter

SRC = Path(__file__).parent / "numbskull spud.md"
OUT = Path(__file__).parent / "numbskull_spud.docx"

DEEP_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
TEAL = RGBColor(0x2D, 0xD4, 0xBF)
SLATE = RGBColor(0x64, 0x74, 0x8B)
BODY_COLOR = RGBColor(0x1A, 0x1A, 0x2E)

# ---------- helpers ----------

def set_font(run, name="Aptos", size=10, color=BODY_COLOR, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_bottom_border(paragraph, color_hex="2DD4BF", width=6):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def latex_to_omml(latex_str):
    """Convert LaTeX string to Office Math XML element."""
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        mathml_tree = ET.fromstring(mathml)
        return mathml_to_omml(mathml_tree)
    except Exception:
        return None


OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def mathml_to_omml(mathml_elem):
    """Convert MathML element tree to OMML. Simplified — handles common cases."""
    # For complex math, fall back to plain text rendering
    # Build an oMath element with the LaTeX as a run
    return None  # We'll use the plain-text fallback


SUB_MAP = str.maketrans("0123456789+-=()aehijklmnoprstuvx",
                         "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
SUP_MAP = str.maketrans("0123456789+-=()niNKk*",
                         "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿⁱᴺᴷᵏ*")


def _extract_brace_arg(s, pos):
    """Extract content of {...} starting at pos (which should point to '{').
    Returns (content, end_pos) where end_pos is after the closing '}'."""
    if pos >= len(s) or s[pos] != '{':
        return ('', pos)
    depth = 0
    start = pos + 1
    for j in range(pos, len(s)):
        if s[j] == '{':
            depth += 1
        elif s[j] == '}':
            depth -= 1
            if depth == 0:
                return (s[start:j], j + 1)
    return (s[start:], len(s))


def _latex_to_unicode(latex_str):
    """Convert LaTeX to readable Unicode text."""
    s = latex_str.strip()

    # --- Pass 1: Handle commands that take {braced args} ---
    # Process from innermost outward by iterating until stable
    for _ in range(5):
        prev = s
        # \frac{a}{b} → (a)/(b)
        m = re.search(r'\\frac\s*\{', s)
        if m:
            num, after_num = _extract_brace_arg(s, m.end() - 1)
            den, after_den = _extract_brace_arg(s, after_num)
            s = s[:m.start()] + '(' + num + ')/(' + den + ')' + s[after_den:]
            continue

        # \sqrt{x} → √(x)
        m = re.search(r'\\sqrt\s*\{', s)
        if m:
            content, end = _extract_brace_arg(s, m.end() - 1)
            s = s[:m.start()] + '\u221A(' + content + ')' + s[end:]
            continue

        # Diacritics: \hat{x} → x̂, \bar{x} → x̄, \tilde{x} → x̃
        for cmd, combining in [('hat', '\u0302'), ('bar', '\u0304'), ('tilde', '\u0303')]:
            m = re.search(r'\\' + cmd + r'\s*\{', s)
            if m:
                content, end = _extract_brace_arg(s, m.end() - 1)
                s = s[:m.start()] + content + combining + s[end:]
                break
        else:
            # Unwrap commands: \text{x} → x, \mathbf{x} → x, etc.
            for cmd in ['text', 'mathbf', 'boldsymbol', 'mathrm', 'mathbb',
                        'operatorname', 'underbrace', 'overbrace', 'mathit',
                        'textbf', 'textit', 'mathcal']:
                m = re.search(r'\\' + cmd + r'\s*\{', s)
                if m:
                    content, end = _extract_brace_arg(s, m.end() - 1)
                    s = s[:m.start()] + content + s[end:]
                    break
            else:
                if s == prev:
                    break

    # --- Pass 2: Subscripts and superscripts with Unicode ---
    # _{...} → subscript chars, ^{...} → superscript chars
    def _sub_replace(m):
        inner = m.group(1)
        converted = inner.translate(SUB_MAP)
        if converted != inner:
            return converted
        return '_(' + inner + ')'

    def _sup_replace(m):
        inner = m.group(1)
        converted = inner.translate(SUP_MAP)
        if converted != inner:
            return converted
        return '^(' + inner + ')'

    s = re.sub(r'_\{([^}]*)\}', _sub_replace, s)
    s = re.sub(r'\^\{([^}]*)\}', _sup_replace, s)

    # Single-char sub/superscripts: _x, ^x (only if followed by a single alphanumeric)
    def _single_sub(m):
        c = m.group(1)
        t = c.translate(SUB_MAP)
        return t if t != c else '_' + c

    def _single_sup(m):
        c = m.group(1)
        t = c.translate(SUP_MAP)
        return t if t != c else '^' + c

    s = re.sub(r'_([0-9a-z])', _single_sub, s)
    s = re.sub(r'\^([0-9a-zA-Z*])', _single_sup, s)

    # --- Pass 3: Symbol replacements ---
    symbols = [
        ('\\alpha', '\u03B1'), ('\\beta', '\u03B2'), ('\\gamma', '\u03B3'),
        ('\\delta', '\u03B4'), ('\\epsilon', '\u03B5'), ('\\varepsilon', '\u03B5'),
        ('\\theta', '\u03B8'), ('\\vartheta', '\u03D1'),
        ('\\lambda', '\u03BB'), ('\\mu', '\u03BC'), ('\\nu', '\u03BD'),
        ('\\pi', '\u03C0'), ('\\rho', '\u03C1'), ('\\sigma', '\u03C3'),
        ('\\tau', '\u03C4'), ('\\phi', '\u03C6'), ('\\varphi', '\u03C6'),
        ('\\psi', '\u03C8'), ('\\omega', '\u03C9'), ('\\chi', '\u03C7'),
        ('\\ell', '\u2113'), ('\\iota', '\u03B9'), ('\\kappa', '\u03BA'),
        ('\\zeta', '\u03B6'), ('\\eta', '\u03B7'), ('\\xi', '\u03BE'),
        ('\\Alpha', '\u0391'), ('\\Gamma', '\u0393'), ('\\Delta', '\u0394'),
        ('\\Phi', '\u03A6'), ('\\Sigma', '\u03A3'), ('\\Omega', '\u03A9'),
        ('\\Theta', '\u0398'), ('\\Lambda', '\u039B'), ('\\Pi', '\u03A0'),
        ('\\Psi', '\u03A8'), ('\\Xi', '\u039E'),
        ('\\infty', '\u221E'), ('\\approx', '\u2248'), ('\\neq', '\u2260'),
        ('\\leq', '\u2264'), ('\\geq', '\u2265'), ('\\sim', '~'),
        ('\\propto', '\u221D'), ('\\equiv', '\u2261'),
        ('\\to', '\u2192'), ('\\rightarrow', '\u2192'), ('\\leftarrow', '\u2190'),
        ('\\Rightarrow', '\u21D2'), ('\\Leftarrow', '\u21D0'),
        ('\\times', '\u00D7'), ('\\cdot', '\u00B7'), ('\\cdots', '\u22EF'),
        ('\\ldots', '\u2026'), ('\\dots', '\u2026'),
        ('\\sum', '\u2211'), ('\\prod', '\u220F'), ('\\int', '\u222B'),
        ('\\partial', '\u2202'), ('\\nabla', '\u2207'),
        ('\\in', '\u2208'), ('\\notin', '\u2209'),
        ('\\subset', '\u2282'), ('\\forall', '\u2200'), ('\\exists', '\u2203'),
        ('\\log', 'log'), ('\\exp', 'exp'), ('\\max', 'max'), ('\\min', 'min'),
        ('\\lim', 'lim'), ('\\sup', 'sup'), ('\\inf', 'inf'),
        ('\\left(', '('), ('\\right)', ')'),
        ('\\left[', '['), ('\\right]', ']'),
        ('\\left|', '|'), ('\\right|', '|'),
        ('\\left\\{', '{'), ('\\right\\}', '}'),
        ('\\{', '{'), ('\\}', '}'),
        ('\\,', '\u2009'), ('\\;', '\u2009'), ('\\!', ''), ('\\quad', '  '),
        ('\\ ', ' '), ('\\mid', ' | '),
    ]
    for old, new in symbols:
        s = s.replace(old, new)

    # --- Pass 4: Cleanup ---
    s = re.sub(r'\\[a-zA-Z]+', '', s)  # remove unknown commands
    s = s.replace('{', '').replace('}', '')  # remove stray braces
    s = re.sub(r'  +', ' ', s).strip()

    return s


def add_math_run(paragraph, latex_str):
    """Add LaTeX as formatted Unicode text."""
    display = _latex_to_unicode(latex_str)
    run = paragraph.add_run(display)
    set_font(run, name="Cambria Math", size=10, italic=True, color=BODY_COLOR)
    return run


def parse_inline(paragraph, text, base_size=10, base_color=BODY_COLOR):
    """Parse inline markdown (bold, italic, code, inline math) and add runs."""
    # Split on inline patterns
    # Order: inline math $...$ , bold **...**, italic *...*, code `...`
    pattern = re.compile(
        r'(\$[^$]+\$)'           # inline math
        r'|(\*\*[^*]+\*\*)'     # bold
        r'|(\*[^*]+\*)'         # italic
        r'|(`[^`]+`)'           # code
    )
    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_font(run, size=base_size, color=base_color)

        if m.group(1):  # inline math
            latex = m.group(1)[1:-1]
            add_math_run(paragraph, latex)
        elif m.group(2):  # bold
            content = m.group(2)[2:-2]
            run = paragraph.add_run(content)
            set_font(run, size=base_size, color=DEEP_BLUE, bold=True)
        elif m.group(3):  # italic
            content = m.group(3)[1:-1]
            run = paragraph.add_run(content)
            set_font(run, size=base_size, color=SLATE, italic=True)
        elif m.group(4):  # code
            content = m.group(4)[1:-1]
            run = paragraph.add_run(content)
            set_font(run, name="Consolas", size=8.5, color=BODY_COLOR)

        pos = m.end()

    # Add remaining plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size, color=base_color)


# ---------- main document builder ----------

def build_doc():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Default paragraph format
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(10)
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.3

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Horizontal rule ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            add_bottom_border(p, "2DD4BF", 4)
            i += 1
            continue

        # --- Display math block ($$...$$) ---
        if line.strip().startswith('$$'):
            math_lines = []
            if line.strip() != '$$':
                # Single-line display math: $$...$$
                inner = line.strip()[2:]
                if inner.endswith('$$'):
                    inner = inner[:-2]
                math_lines.append(inner)
            else:
                # Multi-line display math
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    math_lines.append(lines[i])
                    i += 1
            latex = '\n'.join(math_lines).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_math_run(p, latex)
            i += 1
            continue

        # --- Heading ---
        heading_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            p = doc.add_paragraph()

            if level == 1:
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                # Check if this is the very first heading (title)
                is_title = text.startswith("The Numbskull")
                if is_title:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                run = p.add_run(text)
                set_font(run, name="Aptos Display", size=22 if is_title else 18,
                        color=DEEP_BLUE, bold=True)
                add_bottom_border(p, "2DD4BF", 8 if is_title else 6)

            elif level == 2:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                set_font(run, name="Aptos", size=13, color=DEEP_BLUE, bold=True)

            elif level == 3:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(text)
                set_font(run, name="Aptos", size=11, color=RGBColor(0x33, 0x41, 0x55),
                        bold=True)

            i += 1
            continue

        # --- Table ---
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Collect all table rows
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                table_lines.append(lines[j])
                j += 1

            # Parse header
            headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
            # Skip separator (line 1)
            data_rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.strip('|').split('|')]
                data_rows.append(cells)

            ncols = len(headers)
            table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Style header
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                set_cell_shading(cell, "1E3A5F")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                run = p.add_run(h)
                set_font(run, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)

            # Style data rows
            for ri, row_data in enumerate(data_rows):
                for ci in range(min(len(row_data), ncols)):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    if ri % 2 == 1:
                        set_cell_shading(cell, "F8FAFC")
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.space_before = Pt(1)
                    parse_inline(p, row_data[ci], base_size=8.5)

            # Add small space after table
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

            i = j
            continue

        # --- Blockquote ---
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s*', '', lines[i]))
                i += 1
            quote_text = ' '.join(quote_lines)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="2DD4BF"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # Add shading
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F0FDFA" w:val="clear"/>'
            )
            pPr.append(shading)

            parse_inline(p, quote_text, base_size=9.5)
            continue

        # --- Unordered list ---
        if re.match(r'^[-*]\s+', line.strip()):
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                item_text = re.sub(r'^[-*]\s+', '', lines[i].strip())
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.3)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run("  \u2022  ")
                set_font(run, size=10, color=TEAL, bold=True)
                parse_inline(p, item_text)
                i += 1
            continue

        # --- Ordered list ---
        ol_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if ol_match:
            num = 1
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.3)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(f"  {num}.  ")
                set_font(run, size=10, color=DEEP_BLUE, bold=True)
                parse_inline(p, item_text)
                num += 1
                i += 1
            continue

        # --- Empty line ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph ---
        # Collect continuation lines (non-empty, non-special)
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nl = lines[j]
            if (not nl.strip() or nl.strip().startswith('#') or
                nl.strip().startswith('|') or nl.strip().startswith('$$') or
                nl.strip().startswith('>') or nl.strip().startswith('---') or
                re.match(r'^[-*]\s+', nl.strip()) or
                re.match(r'^\d+\.\s+', nl.strip())):
                break
            para_lines.append(nl)
            j += 1

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Check if this is the subtitle (italic line right after title)
        if full_text.startswith('*') and full_text.endswith('*') and not full_text.startswith('**'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(full_text.strip('*'))
            set_font(run, size=10, color=SLATE, italic=True)
        else:
            parse_inline(p, full_text)

        i = j
        continue

    doc.save(OUT)
    print(f"Written to: {OUT}")
    print(f"Size: {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build_doc()
